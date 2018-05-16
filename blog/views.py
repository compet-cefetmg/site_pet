from django.shortcuts import render, get_object_or_404, redirect, reverse
from django.http import HttpResponse, JsonResponse, HttpResponseForbidden
from .models import Post
from .forms import PostForm
from cefet.models import Pet
from django.contrib.auth.models import User
from django.contrib.auth.decorators import login_required
from datetime import datetime
from django.contrib import messages
from site_pet.settings import MEDIA_ROOT
from django.views.decorators.csrf import csrf_exempt

from docx import *
from docx.shared import Inches,Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import StringIO
import PIL
from PIL import Image
from django.utils.html import strip_tags
from resizeimage import resizeimage
import os

def report(request,id):
    post = get_object_or_404(Post, id=id)
    document = Document()
    stripped = strip_tags(post.text_content)
    document.add_heading(post.title+"\n")
    paragraph = document.add_paragraph("\tDescrição: "+post.text_call)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph = document.add_paragraph("\t"+stripped)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    
    document.add_picture(post.thumbnail,width=Inches(5.0))
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename=relatorio.docx'
    document.save(response)
    return response

def index(request):
    posts = []
    pet_id = request.GET.get('pet', '')
    if pet_id:
        pet = get_object_or_404(Pet, id=pet_id)
        posts = Post.objects.filter(member__in=pet.members.all()).order_by('publish_date').filter(definitive=True)
    else:
        posts = Post.objects.order_by('publish_date').filter(definitive=True)
    context = {'posts': posts, 'name': 'blog.index'}
    return render(request, 'blog/index.html', context)


def post(request, id):
    post = get_object_or_404(Post, id=id)
    posts = Post.objects.filter(definitive=True)
    context = {'post': post, 'name': 'blog.index','posts':posts}
    return render(request, 'blog/post.html', context)


def all(request):
    posts = [(p.id, p.title, p.member.name, p.publish_date.strftime("%d/%m/%y"))
             for p in Post.objects.filter(member__in=request.user.member.pet.members.all()).filter(definitive=True)]
    return JsonResponse({'data': posts}, safe=False)

@login_required
def allInative(request):
    posts = []
    pet_id = request.GET.get('pet', '')
    if pet_id:
        pet = get_object_or_404(Pet, id=pet_id)
        posts = Post.objects.filter(member__in=pet.members.all()).order_by('publish_date').filter(definitive=False)
    else:
        posts = Post.objects.order_by('publish_date').filter(definitive=False)
    context = {'posts': posts, 'name': 'blog.index','prev':True}
    return render(request, 'blog/index.html', context)



@login_required
def add_post(request):
    context = {'title': 'Novo post', 'action': reverse('blog.add_post')}

    if request.method == 'POST':
        form = PostForm(request.POST, request.FILES)

        if form.is_valid():
            post = form.save()
            post.member = request.user.member

            # Saves thumbnail as {{ MEDIA_ROOT }}/blog/{{ post.id }}/thumb
            if post.thumbnail:
                thumbnail_path = os.path.join(MEDIA_ROOT, post.thumbnail.name)
                thumbnail_folder = os.path.join(MEDIA_ROOT, 'blog', str(post.id))
                if not os.path.isdir(thumbnail_folder):
                    os.mkdir(thumbnail_folder)
                post.thumbnail.name = os.path.join('blog', str(post.id), 'thumb')
                os.rename(thumbnail_path, os.path.join(MEDIA_ROOT, post.thumbnail.name))
            post.text_content=post.text_content.replace("<img","<img class='img-responsive'")
            post.save()
            return redirect(reverse('blog.preview', args=(post.id,)))

        context['form'] = form
        return render(request, 'blog/form.html', context, status=400)

    context['form'] = PostForm()
    return render(request, 'blog/form.html', context)


@login_required
def edit_post(request, id,prev):
    post = get_object_or_404(Post, id=id)

    if post.member.pet != request.user.member.pet:
        return HttpResponseForbidden('Você só pode editar posts de seu próprio PET.')

    context = {'title': 'Editar post', 'action': reverse('blog.edit_post', kwargs={'id': id,'prev':0})}

    if request.method == 'POST':
        form = PostForm(request.POST, request.FILES, instance=post)

        if form.is_valid():
            form.save()
            # Removes old thumbnail if needed
            if not post.thumbnail.name and os.path.exists(os.path.join(MEDIA_ROOT, 'blog', str(post.id), 'thumb')):
                os.remove(os.path.join(MEDIA_ROOT, 'blog', str(post.id), 'thumb'))
            if(prev=='1'):
                return redirect(reverse('blog.preview', args=(post.id,)))
            else:
                messages.success(request, 'Post editado com sucesso.')
                return redirect(reverse('blog.index'))

        context['form'] = form
        return render(request, 'blog/form.html', context, status=400)

    context['form'] = PostForm(instance=post)
    return render(request, 'blog/form.html', context)

@login_required
@csrf_exempt# TODO TENTAR FAZER SEM ISSO
def preview(request, id):
    post = get_object_or_404(Post, id=id)
    if(post.definitive==True):
        raise Exception('Already Posted')
    if post.member.pet != request.user.member.pet:
        return HttpResponseForbidden('Você só pode validar posts de seu próprio PET.')
    posts = Post.objects.filter(definitive=True)
    context = {'preview': "true",'post': post, 'name': 'blog.index','posts':posts}
    if 'addposttodb' in request.POST:
        post.definitive=True;
        post.save();
        return redirect(reverse('blog.index'))
    elif 'edtposttodb' in request.POST:
        return redirect(reverse('blog.edit_post', args=(post.id,1,)))
    elif 'remposttodb' in request.POST:
        post.delete()
        return redirect(reverse('blog.index'))
    return render(request, 'blog/post.html', context)
    


@login_required
def delete_post(request):
    post = get_object_or_404(Post, id=request.POST.get('id'))

    if post.member.pet != request.user.member.pet or request.method != 'POST':
        return HttpResponseForbidden('Você só pode remover posts de seu próprio PET.')

    post.delete()

    # Removes thumbnail
    if os.path.exists(os.path.join(MEDIA_ROOT, 'blog', request.POST.get('id'), 'thumb')):
        os.remove(os.path.join(MEDIA_ROOT, 'blog', request.POST.get('id'), 'thumb'))

    return HttpResponse()
